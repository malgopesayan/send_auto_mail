"""
app.py — Job Application Pipeline: backend

What this replaces from the old CLI script:
- Screenshots now live in a Supabase Storage bucket instead of a local
  "screenshots/" folder.
- Results now go into a Supabase table instead of job_applications.xlsx.
- After a screenshot is successfully processed, it is DELETED from
  Supabase Storage (not moved to an "output" folder).
- A web dashboard (static/index.html) shows live progress while the
  pipeline runs, lists all jobs, and includes a chatbot. Asking the
  chatbot to "send mail to X" / "send the TestCo one" actually sends
  the email using GmailService (the exact class you provided) and
  updates that row's Email Status in Supabase.

SETUP
-----
1. pip install -r requirements.txt
2. Run schema.sql in your Supabase project's SQL editor to create the
   job_applications table.
3. In Supabase Storage, create a bucket (default name: "screenshots")
   and upload your recruiter/job screenshots into it.
4. Copy .env.example to .env and fill in:
     SUPABASE_URL, SUPABASE_KEY   (service_role key — needed for storage delete)
     GROQ_API_KEY (and optionally GROQ_API_KEY1..4)
5. Put your Gmail OAuth `credentials.json` next to this file (same as
   before — first send will open a browser to log in once, then it
   reuses token.json).
6. Run:  uvicorn app:app --reload
   Open http://127.0.0.1:8000 in a browser.
"""

import os
import re
import json
import time
import base64
import queue
import mimetypes
import threading
from io import BytesIO
from pathlib import Path
from datetime import date, timedelta

from dotenv import load_dotenv
from PIL import Image, ImageOps
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel

from groq import Groq, RateLimitError, APIStatusError
from supabase import create_client, Client

from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import SupabaseVectorStore
from pypdf import PdfReader
import docx as docx_lib

# Gmail — exact class provided by the user, unchanged.
from gmail_service import GmailService

load_dotenv()

# --------------------------------------------------------------------------
# CONFIG
# --------------------------------------------------------------------------

SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")
SUPABASE_BUCKET = os.environ.get("SUPABASE_BUCKET", "screenshots")
SUPABASE_TABLE = os.environ.get("SUPABASE_TABLE", "job_applications")
# Bucket used to store the candidate's resume that gets attached to every
# outgoing application email. Only ever holds a single file — uploading a
# new one replaces whatever was there before.
SUPABASE_RESUME_BUCKET = os.environ.get("SUPABASE_RESUME_BUCKET", "resumes")

VISION_MODEL = os.environ.get("VISION_MODEL", "qwen/qwen3.8-27b")
CHAT_MODEL = os.environ.get("CHAT_MODEL", "openai/gpt-oss-120b")

# --- RAG / knowledge base (LangChain + Gemini embeddings + pgvector) ---
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
# NOTE: gemini-embedding-2 has a known SDK quirk where passing a *list* of
# strings in one call collapses to a single embedding. We only ever embed
# one document at a time below, so this is safe either way. Switch this env
# var to "models/gemini-embedding-001" if you want the older, more battle-
# tested model instead.
EMBEDDING_MODEL = os.environ.get("GEMINI_EMBEDDING_MODEL", "models/gemini-embedding-2")
EMBEDDING_DIMENSIONS = int(os.environ.get("EMBEDDING_DIMENSIONS", "768"))
KNOWLEDGE_TABLE = os.environ.get("KNOWLEDGE_TABLE", "documents")

API_KEY_ENV_VARS = [
    "GROQ_API_KEY",
    "GROQ_API_KEY1",
    "GROQ_API_KEY2",
    "GROQ_API_KEY3",
    "GROQ_API_KEY4",
]

RETRY_BACKOFF_SECONDS = 3
MAX_IMAGE_DIMENSION = 1568
CROP_WHITESPACE_MARGINS = True

CANDIDATE_NAME = "Sayan Malgope"
CANDIDATE_EMAIL = "malgopesayan19@gmail.com"
CANDIDATE_PHONE = "+91 8670096239"
CANDIDATE_LINKEDIN = "linkedin.com/in/malgopesayan"
CANDIDATE_GITHUB = "github.com/malgopesayan"

TARGET_PROFILE = """
Python Developer / AI Engineer skilled in building scalable backend systems and
APIs using FastAPI, with hands-on experience in Generative AI / LLM deployment
(RAG pipelines, LangChain agent workflows), enterprise authentication (Active
Directory/LDAP), and secure API design. Currently AI Engineer at Catnip Infotech
(Jan 2026 - present); previously Software Developer Intern at Nidhisha
Technologies working on Java/Spring Boot + React.js full-stack development,
REST APIs, Firebase, and AWS deployment. Core stack: Python, FastAPI, Java,
Spring Boot, React.js, SQL/MySQL, LangChain, RAG, scikit-learn, AWS. B.Tech in
Computer Science & Engineering. Looking for backend, full-stack, or AI/GenAI
engineering roles.
""".strip()

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
RESUME_EXTENSIONS = {".pdf", ".doc", ".docx"}


def sanitize_storage_filename(filename: str) -> str:
    """Every resume is stored/attached under one clean, fixed name — e.g.
    'Sayan_Malgope_Resume.pdf' — regardless of what the uploaded file was
    called. Only the original extension is kept."""
    suffix = Path(filename or "").suffix.lower()
    if suffix not in RESUME_EXTENSIONS:
        suffix = ".pdf"
    base = re.sub(r"[^A-Za-z0-9]+", "_", CANDIDATE_NAME).strip("_") or "Resume"
    return f"{base}_Resume{suffix}"

EXTRACTION_PROMPT = f"""
You are looking at a screenshot of a recruiter message, job posting, or
LinkedIn DM. Do two things:

1. Extract job/recruiter details from the image.
2. Draft a short, personalized application email for the candidate below,
   tailored to this specific job posting.

Return ONLY a JSON object (no markdown fences, no commentary):

{{
  "recruiter_email": "email address if visible, else empty string",
  "recruiter_name": "person's name if visible, else empty string",
  "company": "company name if identifiable, else empty string",
  "job_title": "job title / role if identifiable, else empty string",
  "location": "job location if mentioned, else empty string",
  "job_match_score": integer 0-100 rating how well this role matches the
      candidate profile below,
  "priority": one of "High", "Medium", "Low" based on the match score and
      role attractiveness,
  "email_subject": "a short, specific subject line for the application email,
      e.g. 'Application for <Job Title> - {CANDIDATE_NAME}'",
  "email_body": "a concise (roughly 120-180 word) plain-text application
      email body. Address the recruiter by name if known, otherwise use a
      generic greeting. Reference the specific job title/company from this
      screenshot. Highlight 2-3 of the candidate's most relevant skills/
      experience for THIS role (pick from the profile below - don't just
      dump the whole profile). End with a polite call to action and a
      sign-off using the candidate's name and contact details below.
      Do not use markdown formatting - plain text only, ready to paste into
      an email client."
}}

Candidate profile to match against and to draft the email from:
\"\"\"{TARGET_PROFILE}\"\"\"

Candidate contact details to sign the email with:
Name: {CANDIDATE_NAME}
Email: {CANDIDATE_EMAIL}
Phone: {CANDIDATE_PHONE}
LinkedIn: {CANDIDATE_LINKEDIN}
GitHub: {CANDIDATE_GITHUB}

If a field genuinely cannot be determined from the image, use an empty
string for text fields and 0 for job_match_score. Still draft
email_subject and email_body even if some job details are missing -
just keep them a bit more generic in that case.
"""


# --------------------------------------------------------------------------
# CLIENTS
# --------------------------------------------------------------------------

def load_groq_keys() -> list[str]:
    return [os.environ[v] for v in API_KEY_ENV_VARS if os.environ.get(v)]


def get_supabase() -> Client:
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise RuntimeError("SUPABASE_URL / SUPABASE_KEY not set in .env")
    return create_client(SUPABASE_URL, SUPABASE_KEY)


_gmail_service = None


def get_gmail_service() -> GmailService:
    global _gmail_service
    if _gmail_service is None:
        _gmail_service = GmailService()
    return _gmail_service


# --------------------------------------------------------------------------
# RESUME (single file, stored in its own Supabase Storage bucket, attached
# to every outgoing application email)
# --------------------------------------------------------------------------

def get_current_resume_info(supabase: Client) -> dict | None:
    """Returns {"name": ..., "updated_at": ...} for the current resume, or None."""
    files = supabase.storage.from_(SUPABASE_RESUME_BUCKET).list()
    files = [f for f in (files or []) if f.get("name") and not f["name"].startswith(".")]
    if not files:
        return None
    # Only one resume is ever kept, but be defensive and take the most recent.
    files.sort(key=lambda f: f.get("updated_at") or f.get("created_at") or "", reverse=True)
    f = files[0]
    return {"name": f["name"], "updated_at": f.get("updated_at") or f.get("created_at")}


def get_current_resume_bytes(supabase: Client) -> tuple[str, bytes] | tuple[None, None]:
    """Returns (filename, raw_bytes) for the current resume, or (None, None) if none uploaded
    (or if the resumes bucket isn't set up yet — this must never break sending mail)."""
    try:
        info = get_current_resume_info(supabase)
        if not info:
            return None, None
        raw = supabase.storage.from_(SUPABASE_RESUME_BUCKET).download(info["name"])
        return info["name"], raw
    except Exception as exc:  # noqa: BLE001 - resume attach is best-effort, never fatal
        print(f"⚠️  Couldn't fetch resume from '{SUPABASE_RESUME_BUCKET}' bucket, "
              f"sending without an attachment: {exc}")
        return None, None


def clear_resume_bucket(supabase: Client) -> None:
    files = supabase.storage.from_(SUPABASE_RESUME_BUCKET).list()
    names = [f["name"] for f in (files or []) if f.get("name") and not f["name"].startswith(".")]
    if names:
        supabase.storage.from_(SUPABASE_RESUME_BUCKET).remove(names)


# --------------------------------------------------------------------------
# KNOWLEDGE BASE / RAG (LangChain + Gemini embeddings + Supabase pgvector)
#
# Covers every job_applications row + the candidate's resume text, so the
# chatbot can semantically search "everything" instead of only exact-match
# lookups. Indexing is automatic: jobs are (re)indexed on insert/update and
# removed on delete; the resume is (re)indexed on upload.
# --------------------------------------------------------------------------

_embeddings = None


def get_embeddings() -> GoogleGenerativeAIEmbeddings:
    global _embeddings
    if _embeddings is None:
        if not GEMINI_API_KEY:
            raise RuntimeError("GEMINI_API_KEY not set in .env — required for the knowledge-base / RAG chat feature.")
        _embeddings = GoogleGenerativeAIEmbeddings(
            model=EMBEDDING_MODEL,
            google_api_key=GEMINI_API_KEY,
            output_dimensionality=EMBEDDING_DIMENSIONS,
        )
    return _embeddings


def get_vector_store(supabase: Client) -> SupabaseVectorStore:
    return SupabaseVectorStore(
        client=supabase,
        embedding=get_embeddings(),
        table_name=KNOWLEDGE_TABLE,
        query_name="match_documents",
    )


def _delete_indexed_docs(supabase: Client, source_type: str, source_id: str | None = None) -> None:
    q = supabase.table(KNOWLEDGE_TABLE).delete().eq("metadata->>source_type", source_type)
    if source_id is not None:
        q = q.eq("metadata->>source_id", source_id)
    q.execute()


def _job_document_text(job: dict) -> str:
    return "\n".join([
        f"Company: {job.get('company') or '—'}",
        f"Job Title: {job.get('job_title') or '—'}",
        f"Location: {job.get('location') or '—'}",
        f"Recruiter: {job.get('recruiter_name') or '—'} <{job.get('recruiter_email') or '—'}>",
        f"Priority: {job.get('priority') or '—'}",
        f"Match Score: {job.get('job_match_score') or 0}",
        f"Email Status: {job.get('email_status') or '—'}",
        f"Email Subject: {job.get('email_subject') or '—'}",
        f"Email Body: {job.get('email_body') or '—'}",
    ])


def index_job(supabase: Client, job: dict) -> None:
    """Upsert one job's embedding in the knowledge base. Best-effort — a
    failure here (e.g. GEMINI_API_KEY missing) must never break the caller."""
    if not job or job.get("id") is None:
        return
    try:
        job_id = str(job["id"])
        _delete_indexed_docs(supabase, "job", job_id)
        doc = Document(
            page_content=_job_document_text(job),
            metadata={
                "source_type": "job",
                "source_id": job_id,
                "company": job.get("company") or "",
                "job_title": job.get("job_title") or "",
            },
        )
        get_vector_store(supabase).add_documents([doc])
    except Exception as exc:  # noqa: BLE001
        print(f"⚠️  Couldn't index job #{job.get('id')} into the knowledge base: {exc}")


def delete_job_from_index(supabase: Client, job_id: int) -> None:
    try:
        _delete_indexed_docs(supabase, "job", str(job_id))
    except Exception as exc:  # noqa: BLE001
        print(f"⚠️  Couldn't remove job #{job_id} from the knowledge base: {exc}")


def extract_resume_text(filename: str, content: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()
    try:
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(content))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        if suffix == ".docx":
            d = docx_lib.Document(BytesIO(content))
            return "\n".join(p.text for p in d.paragraphs)
    except Exception as exc:  # noqa: BLE001
        print(f"⚠️  Couldn't extract text from resume ({filename}): {exc}")
        return ""
    # Legacy .doc (binary Word format) isn't parseable without extra system
    # tools — skip text extraction for it. PDF/DOCX cover the common case.
    return ""


def index_resume(supabase: Client, filename: str, content: bytes) -> None:
    try:
        _delete_indexed_docs(supabase, "resume")
        text = extract_resume_text(filename, content)
        if not text.strip():
            return
        doc = Document(page_content=text, metadata={"source_type": "resume", "source_id": filename})
        get_vector_store(supabase).add_documents([doc])
    except Exception as exc:  # noqa: BLE001
        print(f"⚠️  Couldn't index resume into the knowledge base: {exc}")


def search_knowledge_base(supabase: Client, query: str, k: int = 6) -> list[Document]:
    return get_vector_store(supabase).similarity_search(query, k=k)


def reindex_all(supabase: Client) -> dict:
    """Bulk (re)index everything — used for the initial backfill of jobs that
    existed before the knowledge base was added, and as a manual fix-up."""
    counts = {"jobs": 0, "resume": 0, "errors": []}

    try:
        supabase.table(KNOWLEDGE_TABLE).delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
    except Exception as exc:  # noqa: BLE001
        counts["errors"].append(f"clear existing index: {exc}")

    resp = supabase.table(SUPABASE_TABLE).select("*").execute()
    for job in resp.data or []:
        try:
            index_job(supabase, job)
            counts["jobs"] += 1
        except Exception as exc:  # noqa: BLE001
            counts["errors"].append(f"job #{job.get('id')}: {exc}")

    try:
        filename, raw = get_current_resume_bytes(supabase)
        if filename and raw:
            index_resume(supabase, filename, raw)
            counts["resume"] = 1
    except Exception as exc:  # noqa: BLE001
        counts["errors"].append(f"resume: {exc}")

    return counts


# --------------------------------------------------------------------------
# IMAGE PREPROCESSING + EXTRACTION (same approach as the CLI version, but
# works on in-memory bytes since images now come from Supabase Storage)
# --------------------------------------------------------------------------

def preprocess_image_bytes(raw_bytes: bytes) -> Image.Image:
    img = Image.open(BytesIO(raw_bytes))
    img = ImageOps.exif_transpose(img)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    if CROP_WHITESPACE_MARGINS:
        gray = img.convert("L")
        diff = ImageOps.invert(gray) if gray.getpixel((0, 0)) > 200 else gray
        bbox = diff.getbbox()
        if bbox:
            pad = 8
            left = max(bbox[0] - pad, 0)
            top = max(bbox[1] - pad, 0)
            right = min(bbox[2] + pad, img.width)
            bottom = min(bbox[3] + pad, img.height)
            img = img.crop((left, top, right, bottom))

    longest_side = max(img.size)
    if longest_side > MAX_IMAGE_DIMENSION:
        scale = MAX_IMAGE_DIMENSION / longest_side
        new_size = (max(1, int(img.width * scale)), max(1, int(img.height * scale)))
        img = img.resize(new_size, Image.LANCZOS)

    return img


def image_bytes_to_data_url(raw_bytes: bytes) -> str:
    img = preprocess_image_bytes(raw_bytes)
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    b64_data = base64.b64encode(buffer.getvalue()).decode("utf-8")
    return f"data:image/png;base64,{b64_data}"


def is_quota_error(exc: Exception) -> bool:
    if isinstance(exc, RateLimitError):
        return True
    if isinstance(exc, APIStatusError):
        return exc.status_code == 429
    return "429" in str(exc) or "rate_limit" in str(exc).lower()


def is_extraction_mostly_empty(fields: dict) -> bool:
    core_fields = (
        fields.get("company", ""),
        fields.get("job_title", ""),
        fields.get("recruiter_name", ""),
        fields.get("recruiter_email", ""),
    )
    return all(not str(v).strip() for v in core_fields)


def extract_fields(client: Groq, raw_bytes: bytes) -> dict:
    data_url = image_bytes_to_data_url(raw_bytes)

    response = client.chat.completions.create(
        model=VISION_MODEL,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": EXTRACTION_PROMPT},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0.3,
        max_completion_tokens=2048,
        response_format={"type": "json_object"},
    )

    raw_text = response.choices[0].message.content or ""
    cleaned = re.sub(r"^```(?:json)?|```$", "", raw_text.strip(), flags=re.MULTILINE).strip()

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return {
            "recruiter_email": "", "recruiter_name": "", "company": "",
            "job_title": "", "location": "", "job_match_score": 0,
            "priority": "", "email_subject": "", "email_body": "",
        }


def extract_with_retries(clients: list[Groq], raw_bytes: bytes, log) -> tuple[dict, bool]:
    """Returns (fields, flagged_low_confidence)."""
    max_retries = max(len(clients), 2)
    low_confidence_fallback = None
    fields = None

    for attempt in range(max_retries):
        client = clients[attempt % len(clients)]
        try:
            candidate = extract_fields(client, raw_bytes)
        except Exception as exc:  # noqa: BLE001
            if is_quota_error(exc):
                log(f"  ⚠️  quota hit (attempt {attempt + 1}/{max_retries}), trying another key...")
            else:
                log(f"  ⚠️  error on attempt {attempt + 1}: {exc.__class__.__name__}: {exc}")
            time.sleep(RETRY_BACKOFF_SECONDS)
            continue

        if is_extraction_mostly_empty(candidate):
            low_confidence_fallback = candidate
            log(f"  ⚠️  extraction came back empty (attempt {attempt + 1}/{max_retries}), retrying...")
            time.sleep(RETRY_BACKOFF_SECONDS)
            continue

        return candidate, False

    if low_confidence_fallback is not None:
        low_confidence_fallback["priority"] = "REVIEW - low OCR confidence"
        return low_confidence_fallback, True

    return None, False


# --------------------------------------------------------------------------
# PIPELINE: Supabase Storage -> Groq extraction -> Supabase table -> delete
# --------------------------------------------------------------------------

pipeline_state = {"running": False}
pipeline_log_queue: "queue.Queue" = queue.Queue()


def run_pipeline():
    pipeline_state["running"] = True

    def log(msg: str):
        pipeline_log_queue.put({"type": "log", "message": msg})

    try:
        groq_keys = load_groq_keys()
        if not groq_keys:
            log("❌ No Groq API keys configured in .env — aborting.")
            pipeline_log_queue.put({"type": "done", "succeeded": 0, "failed": 0})
            return

        clients = [Groq(api_key=k) for k in groq_keys]
        supabase = get_supabase()

        files = supabase.storage.from_(SUPABASE_BUCKET).list()
        image_files = [
            f for f in files
            if Path(f["name"]).suffix.lower() in IMAGE_EXTENSIONS
        ]
        total = len(image_files)
        pipeline_log_queue.put({"type": "start", "total": total})

        if total == 0:
            log("No new screenshots found in the bucket.")
            pipeline_log_queue.put({"type": "done", "succeeded": 0, "failed": 0})
            return

        succeeded, failed = 0, 0

        for idx, file_info in enumerate(image_files, start=1):
            name = file_info["name"]
            log(f"[{idx}/{total}] Processing {name} ...")

            try:
                raw_bytes = supabase.storage.from_(SUPABASE_BUCKET).download(name)
            except Exception as exc:  # noqa: BLE001
                log(f"[{idx}/{total}] ❌ {name}: couldn't download from storage ({exc}). Skipping.")
                failed += 1
                pipeline_log_queue.put({"type": "progress", "idx": idx, "total": total, "status": "failed", "name": name})
                continue

            fields, flagged = extract_with_retries(clients, raw_bytes, log)

            if fields is None:
                log(f"[{idx}/{total}] ❌ {name}: failed on every attempt. Left in the bucket — re-run to retry.")
                failed += 1
                pipeline_log_queue.put({"type": "progress", "idx": idx, "total": total, "status": "failed", "name": name})
                continue

            row = {
                "recruiter_email": fields.get("recruiter_email", ""),
                "recruiter_name": fields.get("recruiter_name", ""),
                "company": fields.get("company", ""),
                "job_title": fields.get("job_title", ""),
                "location": fields.get("location", ""),
                "job_match_score": fields.get("job_match_score", 0) or 0,
                "priority": fields.get("priority", ""),
                "email_subject": fields.get("email_subject", ""),
                "email_body": fields.get("email_body", ""),
                "email_status": "Not Sent",
                "sent_date": None,
                "follow_up_date": None,
                "source_screenshot": name,
            }

            try:
                insert_resp = supabase.table(SUPABASE_TABLE).insert(row).execute()
            except Exception as exc:  # noqa: BLE001
                log(f"[{idx}/{total}] ❌ {name}: extracted fine, but couldn't save to the table "
                    f"({exc}). Left in the bucket so nothing is lost.")
                failed += 1
                pipeline_log_queue.put({"type": "progress", "idx": idx, "total": total, "status": "failed", "name": name})
                continue

            inserted_row = (insert_resp.data or [{}])[0]
            index_job(supabase, {**row, "id": inserted_row.get("id")})

            try:
                supabase.storage.from_(SUPABASE_BUCKET).remove([name])
            except Exception as exc:  # noqa: BLE001
                log(f"[{idx}/{total}] ⚠️  {name}: saved to the table, but couldn't delete from "
                    f"storage ({exc}). You may want to remove it manually.")

            status_icon = "⚠️ " if flagged else "✅"
            log(f"[{idx}/{total}] {status_icon} {name} -> {row['company'] or '?'} | "
                f"{row['job_title'] or '?'} | score={row['job_match_score']} | priority={row['priority']}")
            succeeded += 1
            pipeline_log_queue.put({
                "type": "progress", "idx": idx, "total": total,
                "status": "review" if flagged else "success", "name": name, "row": row,
            })

        log(f"Done. {succeeded} succeeded, {failed} failed.")
        pipeline_log_queue.put({"type": "done", "succeeded": succeeded, "failed": failed})

    except Exception as exc:  # noqa: BLE001 - never let the background thread die silently
        pipeline_log_queue.put({"type": "log", "message": f"❌ Pipeline crashed: {exc}"})
        pipeline_log_queue.put({"type": "done", "succeeded": 0, "failed": 0})
    finally:
        pipeline_state["running"] = False


# --------------------------------------------------------------------------
# CHATBOT — Groq tool-calling over the jobs table + Gmail sending
# --------------------------------------------------------------------------

def find_job_matches(supabase: Client, query: str) -> list[dict]:
    """Used for sending mail — only considers jobs that haven't been sent yet."""
    query = query.strip().lower()
    resp = supabase.table(SUPABASE_TABLE).select("*").neq("email_status", "Sent").execute()
    rows = resp.data or []
    matches = [
        r for r in rows
        if query in (r.get("company") or "").lower()
        or query in (r.get("job_title") or "").lower()
        or query in (r.get("recruiter_name") or "").lower()
        or query in (r.get("recruiter_email") or "").lower()
    ]
    return matches


def find_any_job(supabase: Client, query: str = "", job_id: int | None = None) -> list[dict]:
    """General-purpose lookup across the WHOLE table (any status), for view/update/delete tools."""
    if job_id is not None:
        resp = supabase.table(SUPABASE_TABLE).select("*").eq("id", job_id).limit(1).execute()
        return resp.data or []

    query = (query or "").strip().lower()
    resp = supabase.table(SUPABASE_TABLE).select("*").execute()
    rows = resp.data or []
    if not query:
        return rows
    return [
        r for r in rows
        if query in (r.get("company") or "").lower()
        or query in (r.get("job_title") or "").lower()
        or query in (r.get("recruiter_name") or "").lower()
        or query in (r.get("recruiter_email") or "").lower()
        or query in (r.get("location") or "").lower()
        or query in (r.get("priority") or "").lower()
        or query in (r.get("email_status") or "").lower()
        or query in str(r.get("id"))
    ]


# Columns the chatbot is allowed to read/write on job_applications.
EDITABLE_JOB_COLUMNS = {
    "recruiter_email", "recruiter_name", "company", "job_title", "location",
    "job_match_score", "priority", "email_subject", "email_body",
    "email_status", "sent_date", "follow_up_date",
}


def send_mail_for_job(supabase: Client, row: dict) -> str:
    if not row.get("recruiter_email"):
        return f"'{row.get('company', 'that job')}' has no recruiter email on file, so I can't send it."

    resume_filename, resume_bytes = get_current_resume_bytes(supabase)

    gmail = get_gmail_service()
    gmail.send_mail(
        to_email=row["recruiter_email"],
        subject=row.get("email_subject") or f"Application - {CANDIDATE_NAME}",
        body=row.get("email_body") or "",
        attachment_bytes=resume_bytes,
        attachment_filename=resume_filename,
    )

    today = date.today()
    updates = {
        "email_status": "Sent",
        "sent_date": today.isoformat(),
        "follow_up_date": (today + timedelta(days=7)).isoformat(),
    }
    supabase.table(SUPABASE_TABLE).update(updates).eq("id", row["id"]).execute()
    index_job(supabase, {**row, **updates})

    resume_note = f" (resume attached: {resume_filename})" if resume_filename else " (no resume on file — sent without an attachment)"
    return (f"Sent the application email to {row['recruiter_email']} for "
            f"{row.get('job_title', 'the role')} at {row.get('company', 'that company')}, "
            f"and marked it Sent.{resume_note}")


CHAT_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "send_mail",
            "description": "Send the drafted application email for a specific job to its recruiter, and mark it as Sent. Use this whenever the user asks to send, email, or apply to a specific job/company/recruiter.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Company name, job title, or recruiter name/email to identify which job to send. Use the words the user gave.",
                    }
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_stats",
            "description": "Get counts of jobs by priority and email status. Use for questions like 'how many jobs do I have' or 'how many are high priority'.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "start_processing",
            "description": "Start processing new screenshots from the Supabase bucket into the jobs table. Use when the user asks to process, extract, scan, or run the pipeline on new screenshots.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_jobs",
            "description": "Search or list rows from the job_applications table, across ALL statuses (sent or not). Use for any question about what jobs exist, e.g. 'show me jobs at TestCo', 'list high priority jobs', 'what's the status of the Google application'. Returns up to `limit` matching rows with all columns.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Free-text filter matched against company, job title, recruiter name/email, location, priority, email status, or id. Leave empty to list everything.",
                    },
                    "limit": {
                        "type": "integer",
                        "description": "Max rows to return. Defaults to 20.",
                    },
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_job",
            "description": "Get the full detail of one specific job row, including the full drafted email body. Use when the user wants to see everything about one job, e.g. 'show me the full email for the Infosys one'.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Company, job title, recruiter name/email, or id to identify the job.",
                    }
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "update_job",
            "description": (
                "Edit one or more fields on a job row in job_applications. Use whenever the user asks to change, "
                "correct, or update anything about a job — e.g. fix a recruiter's email, change the priority, "
                "rewrite the email subject/body, change match score, or manually mark a job's email_status. "
                "Editable columns: recruiter_email, recruiter_name, company, job_title, location, job_match_score "
                "(0-100 integer), priority (High/Medium/Low), email_subject, email_body, email_status "
                "(Not Sent/Sent), sent_date (YYYY-MM-DD), follow_up_date (YYYY-MM-DD)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Company, job title, recruiter name/email, or id to identify the job to update.",
                    },
                    "updates": {
                        "type": "object",
                        "description": "Object mapping column name -> new value. Only include the fields being changed.",
                    },
                },
                "required": ["query", "updates"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delete_job",
            "description": "Permanently delete a job row from job_applications. Use when the user asks to remove, delete, or discard a job entry. This cannot be undone, so only call it when the user's intent to delete is clear.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Company, job title, recruiter name/email, or id to identify the job to delete.",
                    }
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_knowledge",
            "description": (
                "Semantic search across the FULL knowledge base: every job application "
                "(company, title, location, recruiter, priority, match score, status, "
                "and the full drafted email) AND the candidate's resume text. Use this "
                "for open-ended, fuzzy, or 'find similar' questions a simple keyword "
                "lookup can't answer — e.g. 'which applications are for backend/Python "
                "roles', 'does my resume mention AWS', 'summarize my high-priority "
                "applications', 'find recruiters who mentioned urgent hiring'. Prefer "
                "list_jobs/get_job for exact lookups by name; use this for anything "
                "broader or content-based, including anything about the resume."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "The question or topic to search for."}
                },
                "required": ["query"],
            },
        },
    },
]


def run_chat_tool(name: str, args: dict, supabase: Client) -> str:
    if name == "send_mail":
        matches = find_job_matches(supabase, args.get("query", ""))
        if not matches:
            return f"I couldn't find an unsent job matching '{args.get('query')}'."
        if len(matches) > 1:
            options = ", ".join(f"{m.get('company')} - {m.get('job_title')}" for m in matches[:5])
            return f"That matches more than one job: {options}. Can you be more specific?"
        return send_mail_for_job(supabase, matches[0])

    if name == "get_stats":
        resp = supabase.table(SUPABASE_TABLE).select("priority,email_status,sent_date").execute()
        rows = resp.data or []
        by_priority = {}
        by_status = {}
        today_str = date.today().isoformat()
        week_ago_str = (date.today() - timedelta(days=7)).isoformat()
        sent_today = 0
        sent_this_week = 0
        for r in rows:
            by_priority[r.get("priority") or "Unknown"] = by_priority.get(r.get("priority") or "Unknown", 0) + 1
            by_status[r.get("email_status") or "Unknown"] = by_status.get(r.get("email_status") or "Unknown", 0) + 1
            sent_date = r.get("sent_date")
            if sent_date == today_str:
                sent_today += 1
            if sent_date and sent_date >= week_ago_str:
                sent_this_week += 1
        return json.dumps({
            "total": len(rows),
            "by_priority": by_priority,
            "by_status": by_status,
            "sent_today": sent_today,
            "sent_last_7_days": sent_this_week,
        })

    if name == "start_processing":
        if pipeline_state["running"]:
            return "The pipeline is already running — check the console panel for progress."
        threading.Thread(target=run_pipeline, daemon=True).start()
        return "Started processing new screenshots. Watch the console panel for live progress."

    if name == "list_jobs":
        limit = args.get("limit") or 20
        rows = find_any_job(supabase, args.get("query", ""))
        rows = rows[: max(1, min(int(limit), 100))]
        if not rows:
            return "No jobs matched that."
        summary = [
            {
                "id": r.get("id"),
                "company": r.get("company"),
                "job_title": r.get("job_title"),
                "location": r.get("location"),
                "priority": r.get("priority"),
                "job_match_score": r.get("job_match_score"),
                "email_status": r.get("email_status"),
                "recruiter_email": r.get("recruiter_email"),
            }
            for r in rows
        ]
        return json.dumps(summary)

    if name in ("get_job", "update_job", "delete_job"):
        query = args.get("query", "")
        job_id = None
        if query.strip().isdigit():
            job_id = int(query.strip())
        matches = find_any_job(supabase, query="" if job_id is not None else query, job_id=job_id)
        if not matches:
            return f"I couldn't find any job matching '{query}'."
        if len(matches) > 1:
            options = ", ".join(f"#{m.get('id')} {m.get('company')} - {m.get('job_title')}" for m in matches[:8])
            return f"That matches more than one job: {options}. Can you be more specific, or give the id?"
        row = matches[0]

        if name == "get_job":
            return json.dumps(row)

        if name == "update_job":
            updates = args.get("updates") or {}
            clean_updates = {k: v for k, v in updates.items() if k in EDITABLE_JOB_COLUMNS}
            rejected = [k for k in updates if k not in EDITABLE_JOB_COLUMNS]
            if not clean_updates:
                return f"None of the fields you gave are editable. Editable fields: {', '.join(sorted(EDITABLE_JOB_COLUMNS))}."
            supabase.table(SUPABASE_TABLE).update(clean_updates).eq("id", row["id"]).execute()
            index_job(supabase, {**row, **clean_updates})
            note = f" (ignored non-editable field(s): {', '.join(rejected)})" if rejected else ""
            return (f"Updated job #{row['id']} ({row.get('company')} - {row.get('job_title')}): "
                    f"set {clean_updates}.{note}")

        if name == "delete_job":
            supabase.table(SUPABASE_TABLE).delete().eq("id", row["id"]).execute()
            delete_job_from_index(supabase, row["id"])
            return f"Deleted job #{row['id']} ({row.get('company')} - {row.get('job_title')})."

    if name == "search_knowledge":
        query = args.get("query", "")
        try:
            docs = search_knowledge_base(supabase, query, k=6)
        except Exception as exc:  # noqa: BLE001
            return (f"Knowledge-base search failed ({exc}). Make sure GEMINI_API_KEY is set "
                    f"and the 'documents' table / match_documents function exist in Supabase.")
        if not docs:
            return "Nothing relevant found in the knowledge base."
        results = [
            {"source": d.metadata.get("source_type"), "ref": d.metadata.get("source_id"), "content": d.page_content}
            for d in docs
        ]
        return json.dumps(results)

    return f"Unknown tool: {name}"


CHAT_SYSTEM_PROMPT = f"""
You are the assistant embedded in {CANDIDATE_NAME}'s job-application pipeline
dashboard. You have full read and write access to the job_applications table
via your tools: you can list/search every job regardless of status, view a
job's full detail (including the drafted email body), edit any editable field
on a job, delete a job outright, report aggregate stats, kick off screenshot
processing, and send application emails (with the candidate's resume
auto-attached if one is on file) on the user's behalf.

You also have search_knowledge — a semantic (RAG) search over EVERYTHING:
every job's full content and the candidate's resume text. Use list_jobs/
get_job for exact lookups by name or id, and search_knowledge for open-ended,
fuzzy, "find similar", or resume-content questions that keyword matching
can't answer well.

For any question about counts — how many total, how many sent, how many
today/this week, breakdowns by priority or status — always use get_stats
first; it already includes sent_today and sent_last_7_days. Only fall back
to list_jobs if get_stats genuinely doesn't cover what was asked.

Be concise and direct. When you change something (send mail, edit a field,
delete a row, start processing), confirm exactly what happened in plain
language. If a request is ambiguous — which job to act on, or which field to
change — ask one short clarifying question instead of guessing. Before
deleting a job, only proceed if the user's intent is clearly a delete request.
"""


class ChatRequest(BaseModel):
    message: str
    history: list[dict] = []


# --------------------------------------------------------------------------
# FASTAPI APP
# --------------------------------------------------------------------------

app = FastAPI(title="Job Application Pipeline")


@app.post("/api/screenshots/upload")
async def api_upload_screenshots(files: list[UploadFile] = File(...)):
    supabase = get_supabase()
    uploaded, skipped, errors = [], [], []

    for f in files:
        suffix = Path(f.filename or "").suffix.lower()
        if suffix not in IMAGE_EXTENSIONS:
            skipped.append(f.filename)
            continue

        content = await f.read()
        mime_type = f.content_type or mimetypes.guess_type(f.filename or "")[0] or "application/octet-stream"
        try:
            supabase.storage.from_(SUPABASE_BUCKET).upload(
                f.filename,
                content,
                file_options={"content-type": mime_type, "upsert": "true"},
            )
            uploaded.append(f.filename)
        except Exception as exc:  # noqa: BLE001
            errors.append({"name": f.filename, "error": str(exc)})

    return {"uploaded": uploaded, "skipped": skipped, "errors": errors}


@app.get("/api/resume")
def api_get_resume():
    supabase = get_supabase()
    try:
        info = get_current_resume_info(supabase)
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=500,
            detail=f"Couldn't reach the '{SUPABASE_RESUME_BUCKET}' bucket — has it been created in Supabase Storage yet? ({exc})",
        )
    return info or {}


@app.post("/api/resume/upload")
async def api_upload_resume(file: UploadFile = File(...)):
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in RESUME_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Resume must be a .pdf, .doc, or .docx file")

    content = await file.read()
    safe_name = sanitize_storage_filename(file.filename or "resume.pdf")
    mime_type = file.content_type or mimetypes.guess_type(safe_name)[0] or "application/octet-stream"

    supabase = get_supabase()
    try:
        # Only one resume is ever kept — clear whatever was there before uploading the new one.
        clear_resume_bucket(supabase)
        supabase.storage.from_(SUPABASE_RESUME_BUCKET).upload(
            safe_name,
            content,
            file_options={"content-type": mime_type, "upsert": "true"},
        )
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=500,
            detail=f"Couldn't upload to the '{SUPABASE_RESUME_BUCKET}' bucket — has it been created in Supabase Storage yet? ({exc})",
        )
    index_resume(supabase, safe_name, content)
    return {"name": safe_name}


@app.delete("/api/resume")
def api_delete_resume():
    supabase = get_supabase()
    try:
        clear_resume_bucket(supabase)
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Couldn't clear the resume bucket: {exc}")
    _delete_indexed_docs(supabase, "resume")
    return {"status": "deleted"}


@app.post("/api/knowledge/reindex-all")
def api_reindex_all():
    supabase = get_supabase()
    try:
        counts = reindex_all(supabase)
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Reindex failed: {exc}")
    return counts


@app.get("/api/jobs")
def api_get_jobs():
    supabase = get_supabase()
    resp = supabase.table(SUPABASE_TABLE).select("*").order("id", desc=True).execute()
    return resp.data or []


@app.post("/api/jobs/{job_id}/send")
def api_send_job(job_id: int):
    supabase = get_supabase()
    resp = supabase.table(SUPABASE_TABLE).select("*").eq("id", job_id).limit(1).execute()
    rows = resp.data or []
    if not rows:
        raise HTTPException(status_code=404, detail="Job not found")
    try:
        message = send_mail_for_job(supabase, rows[0])
    except Exception as exc:  # noqa: BLE001 - never let this crash unhandled
        raise HTTPException(status_code=500, detail=f"Failed to send: {exc}")
    return {"message": message}


@app.post("/api/process/start")
def api_start_process():
    if pipeline_state["running"]:
        return {"status": "already_running"}
    # Drain any stale messages from a previous run before starting fresh.
    while not pipeline_log_queue.empty():
        pipeline_log_queue.get_nowait()
    threading.Thread(target=run_pipeline, daemon=True).start()
    return {"status": "started"}


@app.get("/api/process/stream")
async def api_process_stream():
    def event_gen():
        while True:
            try:
                item = pipeline_log_queue.get(timeout=1)
                yield f"data: {json.dumps(item)}\n\n"
                if item["type"] == "done":
                    break
            except queue.Empty:
                if not pipeline_state["running"]:
                    break
                yield ": keep-alive\n\n"

    return StreamingResponse(event_gen(), media_type="text/event-stream")


@app.post("/api/chat")
def api_chat(req: ChatRequest):
    groq_keys = load_groq_keys()
    if not groq_keys:
        raise HTTPException(status_code=500, detail="No Groq API key configured")
    client = Groq(api_key=groq_keys[0])
    supabase = get_supabase()

    messages = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}]
    messages.extend(req.history)
    messages.append({"role": "user", "content": req.message})

    try:
        reply = None
        for _ in range(5):  # hard cap so a confused model can't loop forever
            response = client.chat.completions.create(
                model=CHAT_MODEL,
                messages=messages,
                tools=CHAT_TOOLS,
                tool_choice="auto",
                temperature=0.4,
            )
            choice = response.choices[0].message

            if not choice.tool_calls:
                reply = choice.content
                break

            messages.append({
                "role": "assistant",
                "content": choice.content,
                "tool_calls": [tc.model_dump() for tc in choice.tool_calls],
            })
            for tool_call in choice.tool_calls:
                try:
                    args = json.loads(tool_call.function.arguments or "{}")
                    result = run_chat_tool(tool_call.function.name, args, supabase)
                except Exception as tool_exc:  # noqa: BLE001 - one bad tool call shouldn't crash the whole reply
                    result = f"That action failed: {tool_exc}"
                messages.append({
                    "role": "tool",
                    "tool_call_id": tool_call.id,
                    "content": result,
                })
        else:
            reply = "That took more steps than expected — could you rephrase or narrow the question?"
    except Exception as exc:  # noqa: BLE001 - never let this crash unhandled
        raise HTTPException(status_code=500, detail=f"Chat failed: {exc}")

    return {"reply": reply}


app.mount("/", StaticFiles(directory="static", html=True), name="static")
